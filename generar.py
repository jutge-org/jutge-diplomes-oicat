from bs4 import BeautifulSoup
from docx import Document
from docx2pdf import convert
from docxcompose.composer import Composer
import os
import requests
import shutil

# pip3 install docx2pdf docxcompose python-docx beautifulsoup4 requests

generate_participation_certs = True
generate_prize_certs = True

contest_name = 'Jutge:FinalOIcat2020'
prize_template = 'DiplomaPlantilla.docx'
participation_template = 'ParticipacioPlantilla.docx'
temp_folder = './cert_temp/'
prize_output_name = 'GeneratedPrizeCertificates'
participation_output_name = 'GeneratedParticipationCertificates'

# Number of certificates
gold = 4 # minim 1
silver = 4
bronze = 4
total_certificates = gold+silver+bronze

# -------------------- Useful functions --------------------

def clear_paragraph(paragraph):
	p_element = paragraph._p
	p_child_elements = [elm for elm in p_element.iterchildren()]
	for child_element in p_child_elements:
		if "'<w:r>'" in str(child_element):
			p_element.remove(child_element)

class Char ():
	def __init__(self,run ,Char : str):
		self.Char = Char
		self.style = run.style
		self.font  = run.font

class Base():
	def __init__(self,Char:Char = None):
		self.style = Char.style
		self.font  = Char.font

class ParagraphHandle ():
	def __init__(self,pagraph):
		self.Chars = []
		self.paragraph = pagraph
		for run in pagraph.runs:
			for x in range(len(run.text)): #convert
				self.Chars.append(Char(run,run.text[x]))

	def replace(self,OldWord,NewWord):
		text = ""
		for x in self.Chars:
			text += x.Char
		fist = text.find(OldWord)

		if fist == -1: return False
		f = Base(self.Chars[fist])

		for i in range(len(OldWord)):
			self.Chars.pop(fist)
		i = 0

		for l in NewWord:
			self.Chars.insert(fist+i,Char(f,l))
			i += 1
		return True

	def build (self):
		if len(self.Chars) == 0 :return
		paraestilo = self.paragraph.style
		clear_paragraph(self.paragraph)
		self.paragraph.style = paraestilo
		runs = []
		fonts = []
		font = self.Chars[0].font
		run = ""
		for x in self.Chars:
			if x.font == font:
				run += x.Char
			else:
				runs.append(run)
				run = x.Char
				fonts.append(font)
				font = x.font
		runs.append(run)
		fonts.append(font)

		for i in range (len(runs)):
			run = self.paragraph.add_run(runs[i])
			fonte = run.font
			fonte.bold = fonts[i].bold
			fonte.color.rgb = fonts[i].color.rgb
			fonte.complex_script = fonts[i].complex_script
			fonte.cs_bold = fonts[i].cs_bold
			fonte.cs_italic = fonts[i].cs_italic
			fonte.double_strike = fonts[i].double_strike
			fonte.emboss = fonts[i].emboss
			fonte.hidden = fonts[i].hidden
			fonte.highlight_color = fonts[i].highlight_color
			fonte.imprint = fonts[i].imprint
			fonte.italic = fonts[i].italic
			fonte.math = fonts[i].math
			fonte.name = fonts[i].name
			fonte.no_proof = fonts[i].no_proof
			fonte.outline = fonts[i].outline
			fonte.rtl = fonts[i].rtl
			fonte.shadow = fonts[i].shadow
			fonte.size = fonts[i].size
			fonte.small_caps = fonts[i].small_caps
			fonte.snap_to_grid = fonts[i].snap_to_grid
			fonte.spec_vanish = fonts[i].spec_vanish
			fonte.strike = fonts[i].strike
			fonte.subscript = fonts[i].subscript
			fonte.superscript = fonts[i].superscript
			fonte.underline = fonts[i].underline
			fonte.web_hidden = fonts[i].web_hidden

def merge_and_convert(cert_name, total_certificates, temp_folder, output_name):
    # Merge the documents
    composer = Composer(Document(temp_folder + cert_name + '0.docx'))
    for cert in [temp_folder + cert_name + str(index) + '.docx' for index in range(1, total_certificates)]:
        composer.append(Document(cert))
    composer.save(output_name + '.docx')

    # Convert to PDF
    convert(output_name + '.docx')

# ----------------------------------------------------------

prize_str = ['MEDALLA D’OR']*gold +\
            ['MEDALLA DE PLATA']*silver +\
            ['MEDALLA DE BRONZE']*bronze
prize_str[0] += ' com a CAMPIÓ ABSOLUT'

# Get contest info
response = requests.get('https://contest.jutge.org/rankings/' + contest_name)
soup = BeautifulSoup(response.text, 'html.parser')

classification = [entry.text.strip() for entry in soup.find_all('td', {'style': 'text-align: left; '})]

if not os.path.exists(temp_folder):
    os.mkdir(temp_folder)

# Generate participation certs
if generate_participation_certs:
    prize_num = 0
    for student in classification:
        document = Document(participation_template)
        for paragraph in document.paragraphs:
            if '[NAME]' in paragraph.text:
                hand = ParagraphHandle(paragraph)
                hand.replace('[NAME]', student)
                hand.build()

        document.save(temp_folder + 'participation' + str(prize_num) +  '.docx')
        prize_num += 1

    merge_and_convert('participation', len(classification), temp_folder, participation_output_name)

# Generate prize certs
if generate_prize_certs:
    prize_num = 0
    classification = classification[:total_certificates]
    for student, prize in zip(classification, prize_str):
        print(student + ': ' + prize)

        document = Document(prize_template)
        for paragraph in document.paragraphs:
            if '[NAME]' in paragraph.text:
                hand = ParagraphHandle(paragraph)
                hand.replace('[NAME]', student)
                hand.build()

            if '[PRIZE]' in paragraph.text:
                hand = ParagraphHandle(paragraph)
                hand.replace('[PRIZE]', prize)
                hand.build()

        document.save(temp_folder + 'prize' + str(prize_num) +  '.docx')
        prize_num += 1

    merge_and_convert('prize', total_certificates, temp_folder, prize_output_name)

# Remofe the temp folder
if os.path.exists(temp_folder):
    shutil.rmtree(temp_folder)
